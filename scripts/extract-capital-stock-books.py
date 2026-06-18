#!/usr/bin/env python3
"""Extract Capital stock book documents into a stock-item import JSON file."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook


SECURITY_BRANCH = "Capital Security Screen"
WINDOWS_BRANCH = "Capital Windows"

DOCX_SOURCES = [
    {
        "path": "/Users/tony_mac/Dropbox Business Dropbox/Tony Commisso/ComGroup/Canberra Stock Book - Screens & Doors - Copy.docx",
        "category": "Screens & Doors",
        "source": "Canberra Stock Book - Screens & Doors",
        "branch": SECURITY_BRANCH,
        "supplier": "Alspec",
    },
    {
        "path": "/Users/tony_mac/Dropbox Business Dropbox/Tony Commisso/ComGroup/Alspec Glass/Canberra Stock Book - Invisi Gard.docx",
        "category": "Invisi-Gard",
        "source": "Canberra Stock Book - Invisi Gard",
        "branch": SECURITY_BRANCH,
        "supplier": "Alspec",
    },
    {
        "path": "/Users/tony_mac/Dropbox Business Dropbox/Tony Commisso/ComGroup/Alspec Glass/Canberra Stock Book - Eco Gard.docx",
        "category": "Eco-Gard",
        "source": "Canberra Stock Book - Eco Gard",
        "branch": SECURITY_BRANCH,
        "supplier": "Alspec",
    },
]

XLSX_SOURCES = [
    {
        "path": "/Users/tony_mac/Dropbox Business Dropbox/Tony Commisso/ComGroup/Capital Windows.xlsx",
        "category": "Capital Windows",
        "source": "Capital Windows",
        "branch": WINDOWS_BRANCH,
    },
]


def clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\xa0", " ")).strip()


def stable_hash(value: str, length: int = 10) -> str:
    return hashlib.sha1(value.encode("utf-8")).hexdigest()[:length].upper()


def short_code(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9-]+", "-", value.upper()).strip("-")
    return value[:50] or "ITEM"


def code_with_colour(item_code: str, colour_code: str) -> str:
    item_code = short_code(item_code)
    colour_code = short_code(colour_code)
    if not colour_code:
        return item_code
    combined = f"{item_code}-{colour_code}"
    if len(combined) <= 50:
        return combined
    return f"{item_code[:36]}-{stable_hash(combined, 12)}"


def generated_code(prefix: str, *parts: str) -> str:
    return f"{prefix}-{stable_hash('|'.join(parts), 12)}"


def parse_length_metres(text: str) -> float | None:
    matches = re.findall(r"(?:x|X)\s*(\d+(?:\.\d+)?)\s*m\b", text)
    if not matches:
        return None
    try:
        return float(matches[-1])
    except ValueError:
        return None


def is_colour_code(value: str) -> bool:
    if not value or len(value) > 10:
        return False
    if " " in value:
        return False
    return bool(re.fullmatch(r"[A-Za-z0-9-]+", value))


def dedupe(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        key = value.casefold()
        if value and key not in seen:
            seen.add(key)
            result.append(value)
    return result


def colour_from_cells(cells: list[str], item_description: str) -> tuple[str, str]:
    remaining = dedupe([
        value
        for value in cells[2:]
        if value and value.casefold() != item_description.casefold()
    ])
    colour_code = ""
    for value in remaining:
        if is_colour_code(value):
            colour_code = value
            break
    descriptions = [
        value
        for value in remaining
        if value.casefold() != colour_code.casefold()
        and value.casefold() != item_description.casefold()
    ]
    colour_description = max(descriptions, key=len, default="")
    return colour_code, colour_description


def description_for(item: dict[str, Any]) -> str:
    lines = [
        f"Source: {', '.join(item['sources'])}",
        f"Original item code: {item['rawCode']}",
    ]
    if item.get("colourCode"):
        lines.append(f"Colour code: {item['colourCode']}")
    if item.get("colour"):
        lines.append(f"Colour: {item['colour']}")
    if item.get("sourceAmount"):
        lines.append(f"Source amount/unit: {item['sourceAmount']}")
    return "\n".join(lines)


def merge_item(items: dict[tuple[str, str], dict[str, Any]], item: dict[str, Any]) -> None:
    key = (item["branch"], item["code"])
    existing = items.get(key)
    if not existing:
        item["sources"] = [item["source"]]
        item["categories"] = [item["category"]]
        item["description"] = description_for(item)
        items[key] = item
        return

    if item["source"] not in existing["sources"]:
        existing["sources"].append(item["source"])
    if item["category"] not in existing["categories"]:
        existing["categories"].append(item["category"])
    if not existing.get("supplier") and item.get("supplier"):
        existing["supplier"] = item["supplier"]
    if not existing.get("sourceFullLength") and item.get("sourceFullLength"):
        existing["sourceFullLength"] = item["sourceFullLength"]
    if not existing.get("colour") and item.get("colour"):
        existing["colour"] = item["colour"]
    existing["category"] = " / ".join(existing["categories"])[:100]
    existing["description"] = description_for(existing)


def extract_docx(source: dict[str, Any], items: dict[tuple[str, str], dict[str, Any]]) -> dict[str, int]:
    path = Path(source["path"])
    if not path.exists():
        raise FileNotFoundError(path)

    document = Document(path)
    scanned = 0
    imported = 0
    for table in document.tables:
        rows = table.rows
        if len(rows) < 2:
            continue
        for row in rows[1:]:
            cells = [clean(cell.text) for cell in row.cells]
            scanned += 1
            item_code = cells[0] if len(cells) > 0 else ""
            item_description = cells[1] if len(cells) > 1 else ""
            if not item_code or not item_description:
                continue
            if item_code.casefold() == "item code" or item_description.casefold() == "item description":
                continue

            colour_code, colour_description = colour_from_cells(cells, item_description)
            code = code_with_colour(item_code, colour_code)
            item = {
                "branch": source["branch"],
                "source": source["source"],
                "category": source["category"],
                "code": code,
                "rawCode": item_code,
                "name": item_description[:255],
                "serialNumber": None,
                "unit": "EA",
                "unitType": "unit",
                "conditionIndicator": "new",
                "sourceFullLength": parse_length_metres(item_description),
                "actualSize": None,
                "supplier": source.get("supplier"),
                "costPrice": None,
                "colourCode": colour_code,
                "colour": colour_description,
                "sourceAmount": None,
            }
            merge_item(items, item)
            imported += 1
    return {"tables": len(document.tables), "scannedRows": scanned, "importableRows": imported}


def find_header(sheet) -> tuple[int, dict[str, int]]:
    for row_number in range(1, min(sheet.max_row, 20) + 1):
        values = [clean(sheet.cell(row=row_number, column=col).value).casefold() for col in range(1, sheet.max_column + 1)]
        if "supplier" in values and "description" in values:
            return row_number, {value: index + 1 for index, value in enumerate(values) if value}
    raise ValueError(f"Could not find Supplier/Description header in {sheet.title}")


def normal_unit(amount: str, description: str) -> str:
    lowered = amount.casefold()
    if lowered in {"roll", "ea", "each", "unit", "pair", "box"}:
        return "ROLL" if lowered == "roll" else "EA"
    if "roll" in description.casefold():
        return "ROLL"
    return "EA"


def extract_xlsx(source: dict[str, Any], items: dict[tuple[str, str], dict[str, Any]]) -> dict[str, int]:
    path = Path(source["path"])
    if not path.exists():
        raise FileNotFoundError(path)

    workbook = load_workbook(path, data_only=True)
    scanned = 0
    imported = 0
    for sheet in workbook.worksheets:
        header_row, columns = find_header(sheet)
        for row_number in range(header_row + 1, sheet.max_row + 1):
            scanned += 1
            supplier = clean(sheet.cell(row=row_number, column=columns.get("supplier", 1)).value)
            raw_code = clean(sheet.cell(row=row_number, column=columns.get("code", 2)).value)
            name = clean(sheet.cell(row=row_number, column=columns.get("description", 3)).value)
            amount = clean(sheet.cell(row=row_number, column=columns.get("amount", 4)).value)
            if not name:
                continue
            code = short_code(raw_code) if raw_code else generated_code("CW", supplier, name, amount)
            item = {
                "branch": source["branch"],
                "source": f"{source['source']} - {sheet.title}",
                "category": source["category"],
                "code": code,
                "rawCode": raw_code or code,
                "name": name[:255],
                "serialNumber": None,
                "unit": normal_unit(amount, name),
                "unitType": "unit",
                "conditionIndicator": "new",
                "sourceFullLength": parse_length_metres(name),
                "actualSize": None,
                "supplier": supplier or None,
                "costPrice": None,
                "colourCode": None,
                "colour": None,
                "sourceAmount": amount or None,
            }
            merge_item(items, item)
            imported += 1
    return {"worksheets": len(workbook.worksheets), "scannedRows": scanned, "importableRows": imported}


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract Capital stock books to JSON.")
    parser.add_argument("--output", default="/tmp/capital-stock-import.json")
    args = parser.parse_args()

    items: dict[tuple[str, str], dict[str, Any]] = {}
    sources: list[dict[str, Any]] = []
    for source in DOCX_SOURCES:
        summary = extract_docx(source, items)
        sources.append({**source, "summary": summary})
    for source in XLSX_SOURCES:
        summary = extract_xlsx(source, items)
        sources.append({**source, "summary": summary})

    output_items = list(items.values())
    for item in output_items:
        item["category"] = " / ".join(item.get("categories") or [item["category"]])[:100]
        item["description"] = description_for(item)
        item.pop("categories", None)

    by_branch: dict[str, int] = {}
    for item in output_items:
        by_branch[item["branch"]] = by_branch.get(item["branch"], 0) + 1

    payload = {
        "sources": sources,
        "summary": {
            "items": len(output_items),
            "byBranch": by_branch,
        },
        "items": output_items,
    }
    output_path = Path(args.output)
    output_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    print(json.dumps(payload["summary"], indent=2))
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
